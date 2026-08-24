"""
Tests for the document generator (Milestone 5).

Three properties matter most here, each corresponding to a real bug
found either in this project's earlier desktop-app version or caught
again while porting the logic to this backend:

1. Bold preservation across runs — placeholder substitution must never
   collapse a paragraph's runs, which would silently destroy bold on a
   specific word (a name, a CTC amount) while leaving the rest correct.
2. CTC table formatting must exactly match measured real-letter values
   (column widths, border color, alignment, two-line header, spacer
   position) — not approximations.
3. Indian digit grouping (12,00,000) in the CTC table specifically —
   caught by an actual LibreOffice render during this milestone: the
   table cells were using Python's default Western comma grouping
   (1,200,000) even though the text-placeholder formatter already had
   the correct Indian grouping. The bug was in a second, easy-to-forget
   call site, not the shared formatter itself.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.services.ctc_engine import LineItemInput, evaluate_structure
from app.services.document_generator import MandatoryPlaceholderError, generate_docx


def test_bold_is_preserved_per_run_not_collapsed() -> None:
    blocks = [
        {"type": "paragraph", "runs": [
            {"text": "Dear "},
            {"text": "{{employee_name}}", "bold": True},
            {"text": ", welcome to "},
            {"text": "{{company_name}}", "bold": True},
            {"text": "."},
        ]},
    ]
    placeholder_data = {"employee_name": "Avinash Singh", "company_name": "iBridge Techsoft"}
    docx_bytes = generate_docx(blocks, placeholder_data, mandatory_placeholders=[])
    doc = Document(BytesIO(docx_bytes))
    para = doc.paragraphs[0]

    assert para.text == "Dear Avinash Singh, welcome to iBridge Techsoft."
    name_run = next(r for r in para.runs if "Avinash" in r.text)
    assert name_run.bold is True
    company_run = next(r for r in para.runs if "iBridge Techsoft" in r.text)
    assert company_run.bold is True
    plain_run = next(r for r in para.runs if r.text == "Dear ")
    assert plain_run.bold in (False, None)


def _sample_ctc_line_items() -> list[LineItemInput]:
    return [
        LineItemInput(key="basic_monthly", label="Basic Salary", section="Earnings", order=1,
                       guided_type="percent_of", guided_params={"base": "monthly_ctc", "percent": 50}),
        LineItemInput(key="hra_monthly", label="HRA", section="Earnings", order=2,
                       guided_type="percent_of", guided_params={"base": "basic_monthly", "percent": 40}),
        LineItemInput(key="total_earnings_monthly", label="Total Earnings (A)", section="Earnings", order=3,
                       formula="basic_monthly + hra_monthly", is_subtotal=True),
    ]


def test_ctc_table_uses_measured_column_widths() -> None:
    computed = evaluate_structure(_sample_ctc_line_items(), annual_ctc=1200000)
    blocks = [{"type": "ctcTable"}]
    docx_bytes = generate_docx(blocks, {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    table = doc.tables[0]
    assert tuple(c.width for c in table.columns) == (2120900, 1155700, 1155700)


def test_ctc_table_border_color_is_auto_not_gray() -> None:
    computed = evaluate_structure(_sample_ctc_line_items(), annual_ctc=1200000)
    docx_bytes = generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    tbl_pr = doc.tables[0]._tbl.tblPr
    top_border = tbl_pr.find(qn('w:tblBorders')).find(qn('w:top'))
    assert top_border.get(qn('w:color')) == "auto"


def test_ctc_table_numeric_columns_are_centered() -> None:
    computed = evaluate_structure(_sample_ctc_line_items(), annual_ctc=1200000)
    docx_bytes = generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    numeric_cell = doc.tables[0].rows[1].cells[1]
    assert numeric_cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_ctc_table_header_has_real_two_line_break() -> None:
    """Regression test for the desktop-app bug: `run.text = second_line`
    after `add_break()` silently strips the break, collapsing 'Yearly'
    and '(INR)' onto one line."""
    computed = evaluate_structure(_sample_ctc_line_items(), annual_ctc=1200000)
    docx_bytes = generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    header_cell = doc.tables[0].rows[0].cells[1]
    full_text = "".join(r.text for r in header_cell.paragraphs[0].runs)
    # The break renders as a literal "\n" within the run's own .text when
    # read back via python-docx — that's correct (confirms the break
    # element is genuinely present), not something to strip out here.
    assert "\n(INR)" in full_text, f"Expected a real line break before '(INR)', got: {full_text!r}"
    assert full_text == "Yearly\n(INR)"


def test_ctc_table_uses_indian_digit_grouping_not_western() -> None:
    """
    Regression test for the exact bug found via LibreOffice visual
    rendering during this milestone: the table cells used Python's
    default `f"{n:,.0f}"` (Western grouping, "6,00,000" would render as
    "600,000") instead of the shared `format_inr` helper (Indian
    grouping) — a second, easy-to-forget call site even though the
    correct formatter already existed and was used elsewhere (text
    placeholders).
    """
    computed = evaluate_structure(_sample_ctc_line_items(), annual_ctc=1200000)
    docx_bytes = generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    basic_row_cells = [c.text for c in doc.tables[0].rows[1].cells]
    # Basic = 50% of monthly_ctc; annual_ctc=1,200,000 -> yearly Basic = 6,00,000
    assert "6,00,000" in basic_row_cells, f"Expected Indian grouping, got: {basic_row_cells}"
    assert "600,000" not in basic_row_cells, "Found Western grouping — the exact bug this test guards against"


def test_ctc_table_spacer_row_after_subtotal_not_before() -> None:
    items = [
        LineItemInput(key="a", label="Item A", section="Earnings", order=1, formula="monthly_ctc * 0.5"),
        LineItemInput(key="b", label="Total (subtotal)", section="Earnings", order=2,
                       formula="a", is_subtotal=True, spacer_after=True),
        LineItemInput(key="c", label="Item C", section="Earnings", order=3, formula="monthly_ctc * 0.1"),
    ]
    computed = evaluate_structure(items, annual_ctc=1200000)
    docx_bytes = generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=computed)
    doc = Document(BytesIO(docx_bytes))
    row_texts = [row.cells[0].text for row in doc.tables[0].rows]
    # Expect: header, Item A, Total (subtotal), <blank>, Item C — spacer AFTER the subtotal
    total_index = row_texts.index("Total (subtotal)")
    assert row_texts[total_index + 1] == "", "Spacer row must come immediately after the subtotal"
    assert row_texts[total_index + 2] == "Item C"


def test_mandatory_placeholder_blocks_generation() -> None:
    blocks = [{"type": "paragraph", "runs": [{"text": "Reason: {{reason}}"}]}]
    with pytest.raises(MandatoryPlaceholderError, match="reason"):
        generate_docx(blocks, {"reason": ""}, mandatory_placeholders=["reason"])


def test_mandatory_placeholder_with_value_does_not_block() -> None:
    blocks = [{"type": "paragraph", "runs": [{"text": "Reason: {{reason}}"}]}]
    docx_bytes = generate_docx(blocks, {"reason": "Excellent performance"}, mandatory_placeholders=["reason"])
    doc = Document(BytesIO(docx_bytes))
    assert "Excellent performance" in doc.paragraphs[0].text


def test_ctc_table_without_structure_blocks_generation() -> None:
    with pytest.raises(MandatoryPlaceholderError, match="CTC"):
        generate_docx([{"type": "ctcTable"}], {}, mandatory_placeholders=[], ctc_computed_items=None)


def test_no_unresolved_placeholders_in_output() -> None:
    blocks = [
        {"type": "paragraph", "runs": [{"text": "Dear {{employee_name}}, your role is {{designation}}."}]},
    ]
    docx_bytes = generate_docx(
        blocks, {"employee_name": "Priya", "designation": "Engineer"}, mandatory_placeholders=[]
    )
    doc = Document(BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "{{" not in full_text


def test_generation_blocked_before_any_file_is_produced() -> None:
    """A blocked generation must raise before touching python-docx at
    all — never produce a half-written file."""
    blocks = [{"type": "paragraph", "runs": [{"text": "{{mandatory_field}}"}]}]
    with pytest.raises(MandatoryPlaceholderError):
        generate_docx(blocks, {}, mandatory_placeholders=["mandatory_field"])
