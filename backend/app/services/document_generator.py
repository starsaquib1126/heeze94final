"""
Document Generator — compiles block-based letter content into a real
.docx, substituting placeholders and injecting the CTC breakup table
and signature image.

Two pieces of logic are ported EXACTLY from the desktop app, because
both were the product of real bugs found through actual visual
verification (rendering with LibreOffice and comparing against the
real iBridge letters pixel-by-pixel), not just written once and assumed
correct:

  1. Placeholder substitution happens per-run, never by collapsing a
     paragraph's runs into one. Collapsing silently destroys bold/italic
     on specific words (a bolded name, a bolded CTC amount) — this was a
     real, shipped bug in an earlier version of this letter-generation
     logic.

  2. The CTC table's exact formatting — column widths (measured in EMU
     from a real iBridge letter, not guessed), border color ("auto"/
     black, not an approximated gray), centered (not right-aligned)
     numeric columns, a genuine two-line "Yearly / (INR)" header built
     with `run.add_break()` + `run.add_text()` (using `run.text = ...`
     after a break silently strips the break — a real bug caught by
     rendering the output and comparing it to the original), and a
     blank spacer row placed AFTER a subtotal/spacer_after row rather
     than before it (also caught by visual comparison, not code review).
"""

from __future__ import annotations

import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt
from io import BytesIO

from app.services.ctc_engine import ComputedLineItem
from app.services.placeholder_resolver import format_inr

_PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

# Measured directly from a real iBridge Offer/Appointment letter's Schedule A /
# Annexure A table — not guessed round numbers.
_CTC_COL_WIDTHS_EMU = (2120900, 1155700, 1155700)  # label, yearly, monthly
_CTC_BORDER_COLOR = "auto"   # Word's "Automatic" (black) — not an approximated gray
_CTC_BORDER_SIZE = "4"       # 0.5pt


class MandatoryPlaceholderError(ValueError):
    """Raised when a template's mandatory placeholder has no resolved value."""


def render_text(text: str, data: dict[str, str]) -> str:
    def _sub(match: re.Match) -> str:
        return str(data.get(match.group(1), ""))
    return _PLACEHOLDER_PATTERN.sub(_sub, text)


def _apply_run(paragraph, run_data: dict, placeholder_data: dict[str, str]) -> None:
    text = render_text(run_data.get("text", ""), placeholder_data)
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = run_data.get("bold", False)
    run.italic = run_data.get("italic", False)
    run.underline = run_data.get("underline", False)


def _apply_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = tbl_pr.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): _CTC_BORDER_SIZE, qn("w:space"): "0",
            qn("w:color"): _CTC_BORDER_COLOR,
        })
        borders.append(border_el)
    tbl_pr.append(borders)
    layout = tbl_pr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
    tbl_pr.append(layout)


def _set_row_cell_widths(row) -> None:
    for cell, width_emu in zip(row.cells, _CTC_COL_WIDTHS_EMU):
        cell.width = Emu(width_emu)


def _prevent_row_split(row) -> None:
    """Stops this single row from being cut across a page boundary —
    a defensive second layer alongside the page-break-before-table in
    generate_docx, in case a future, longer structure still doesn't
    fully fit even when starting fresh at the top of a page."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.makeelement(qn("w:cantSplit"), {})
    tr_pr.append(cant_split)


def _set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_header_cell_two_line(cell, first_line: str, second_line: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run(first_line)
    run1.bold = True
    run1.font.name = "Calibri"
    run2 = paragraph.add_run()
    run2.bold = True
    run2.font.name = "Calibri"
    run2.add_break()
    # add_text() APPENDS — setting run2.text after add_break() would
    # silently strip the break just added (a real bug caught by rendering
    # the output: "Yearly" and "(INR)" collapsed onto one line).
    run2.add_text(second_line)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_ctc_table(doc: Document, computed_items: list[ComputedLineItem]):
    table = doc.add_table(rows=0, cols=3)
    _apply_table_borders(table)
    for col, width_emu in zip(table.columns, _CTC_COL_WIDTHS_EMU):
        col.width = Emu(width_emu)

    current_section = None
    for index, item in enumerate(computed_items):
        if item.section != current_section:
            current_section = item.section
            header_row = table.add_row()
            _set_row_cell_widths(header_row)
            _prevent_row_split(header_row)
            _set_cell_text(header_row.cells[0], current_section, bold=True)
            _set_header_cell_two_line(header_row.cells[1], "Yearly", "(INR)")
            _set_header_cell_two_line(header_row.cells[2], "Monthly", "(INR)")

        row = table.add_row()
        _set_row_cell_widths(row)
        _prevent_row_split(row)
        yearly_text = item.yearly if isinstance(item.yearly, str) else format_inr(item.yearly)
        monthly_text = item.monthly if isinstance(item.monthly, str) else format_inr(item.monthly)
        _set_cell_text(row.cells[0], item.label, bold=item.is_subtotal)
        _set_cell_text(row.cells[1], yearly_text, bold=item.is_subtotal, center=True)
        _set_cell_text(row.cells[2], monthly_text, bold=item.is_subtotal, center=True)

        is_last_item = index == len(computed_items) - 1
        if item.spacer_after and not is_last_item:
            # Blank row comes AFTER a total-like row (visual separation from
            # what follows), not before it — and never after the very last
            # row, since there's nothing left to separate it from.
            spacer = table.add_row()
            _set_row_cell_widths(spacer)
            _prevent_row_split(spacer)
            for c in spacer.cells:
                _set_cell_text(c, "")

    return table


def _add_signature_image(doc: Document, signature_bytes: bytes | None) -> None:
    paragraph = doc.add_paragraph()
    if signature_bytes:
        run = paragraph.add_run()
        run.add_picture(BytesIO(signature_bytes), width=Inches(1.5))
    else:
        paragraph.add_run("[Authorized Signature]").italic = True


def _add_logo_to_header(doc: Document, logo_bytes: bytes | None) -> None:
    if not logo_bytes:
        return
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(BytesIO(logo_bytes), width=Inches(1.2))


def extract_placeholders(blocks: list[dict]) -> set[str]:
    """Same scanning logic as letter_template_service.extract_placeholders —
    duplicated here deliberately rather than imported, since this module
    must never depend on the database-touching service layer (keeps this
    module trivially unit-testable with plain dicts, no mocks required)."""
    found: set[str] = set()

    def scan_text(text: str) -> None:
        found.update(_PLACEHOLDER_PATTERN.findall(text))

    for block in blocks:
        block_type = block.get("type")
        if block_type in ("paragraph", "heading"):
            for run in block.get("runs", []):
                scan_text(run.get("text", ""))
        elif block_type in ("bulletList", "numberedList"):
            for item in block.get("items", []):
                for run in item.get("runs", []):
                    scan_text(run.get("text", ""))
        elif block_type == "ctcTable":
            found.add("ctc_breakup_table")

    return found


def generate_docx(
    blocks: list[dict],
    placeholder_data: dict[str, str],
    mandatory_placeholders: list[str],
    ctc_computed_items: list[ComputedLineItem] | None = None,
    logo_bytes: bytes | None = None,
    signature_bytes: bytes | None = None,
    letter_type: str | None = None,
) -> bytes:
    """
    Compile block content into a .docx and return its raw bytes.

    Raises MandatoryPlaceholderError if any placeholder flagged mandatory
    on the template has no resolved value — this is checked BEFORE any
    document content is written, so a blocked generation never produces
    a half-finished file.
    """
    used_placeholders = extract_placeholders(blocks)
    missing_mandatory = [
        p for p in mandatory_placeholders
        if p in used_placeholders and not placeholder_data.get(p)
    ]
    if missing_mandatory:
        raise MandatoryPlaceholderError(
            f"These required fields are missing: {', '.join(missing_mandatory)}. "
            f"Generation blocked to avoid a letter with blank fields."
        )
    if "ctc_breakup_table" in used_placeholders and ctc_computed_items is None:
        raise MandatoryPlaceholderError(
            "This template includes a CTC breakup table but no CTC structure was provided."
        )

    doc = Document()

    # Explicit document-wide defaults, matching the real uploaded
    # templates exactly: Carlito (not Calibri — visually near-identical
    # but a genuinely different font name, and the real documents
    # consistently use Carlito), 11pt.
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Carlito"
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(8)
    normal_style.paragraph_format.line_spacing = 1.15

    # Margins genuinely differ per letter type in the real templates —
    # not a single uniform default. Falls back to a reasonable default
    # for any letter_type not in this table.
    MARGINS_BY_LETTER_TYPE = {
        "offer": (1.31, 1.40, 0.68, 1.01),        # top, bottom, left, right (inches)
        "appointment": (0.94, 1.42, 0.85, 0.75),
        "hike": (0.68, 0.74, 0.75, 1.11),
        "relieving": (0.65, 0.19, 0.57, 1.18),
    }
    top, bottom, left, right = MARGINS_BY_LETTER_TYPE.get(letter_type or "", (1.0, 1.0, 1.0, 1.0))
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    _add_logo_to_header(doc, logo_bytes)

    for i, block in enumerate(blocks):
        block_type = block.get("type")
        next_block_is_table = i + 1 < len(blocks) and blocks[i + 1].get("type") == "ctcTable"

        if block_type == "heading":
            if next_block_is_table:
                doc.add_page_break()
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_data in block.get("runs", []):
                _apply_run(paragraph, {**run_data, "bold": True}, placeholder_data)
                # The real "Title" style this matches (based_on Normal,
                # centered, bold, single underline) does NOT override
                # font size — it's the same 11pt as body text, not a
                # larger heading size.
                paragraph.runs[-1].underline = run_data.get("underline", True)

        elif block_type == "paragraph":
            if next_block_is_table:
                doc.add_page_break()
            paragraph = doc.add_paragraph()
            align = block.get("align")
            if align == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run_data in block.get("runs", []):
                _apply_run(paragraph, run_data, placeholder_data)

        elif block_type == "bulletList":
            for item in block.get("items", []):
                paragraph = doc.add_paragraph(style="List Bullet")
                for run_data in item.get("runs", []):
                    _apply_run(paragraph, run_data, placeholder_data)

        elif block_type == "numberedList":
            for item in block.get("items", []):
                paragraph = doc.add_paragraph(style="List Number")
                for run_data in item.get("runs", []):
                    _apply_run(paragraph, run_data, placeholder_data)

        elif block_type == "ctcTable":
            # An explicit page break always breaks at exactly the point
            # it's inserted, regardless of any keep_with_next setting on
            # a paragraph before it (that property only influences
            # automatic breaks the layout engine makes on its own, not
            # where a manual break lands) — so a page break placed here,
            # after a label like "Annexure A" or "Your CTC breakup is as
            # follows:" was already rendered, permanently strands that
            # label on the previous page no matter what formatting is
            # applied to it after the fact.
            #
            # The fix has to happen earlier: whichever paragraph/heading
            # block immediately precedes this one inserts the page break
            # itself, BEFORE it renders its own text (see
            # next_block_is_table above) — so the label and the table
            # both start together on the fresh page. If nothing precedes
            # the table at all, break here as a fallback so the table
            # still gets a full page to itself.
            if i == 0 or blocks[i - 1].get("type") not in ("paragraph", "heading"):
                doc.add_page_break()
            build_ctc_table(doc, ctc_computed_items or [])

        elif block_type == "signature":
            _add_signature_image(doc, signature_bytes)

        elif block_type == "spacer":
            doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
